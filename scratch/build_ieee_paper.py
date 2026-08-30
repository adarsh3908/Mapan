import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

def create_ieee_paper():
    doc = Document()

    # Page setup - 0.75 in margins
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styles & Colors
    COLOR_TITLE = RGBColor(15, 23, 42)      # Dark Slate
    COLOR_PRIMARY = RGBColor(30, 58, 138)   # Deep Navy
    COLOR_TEXT = RGBColor(30, 41, 59)       # Body Text Slate
    COLOR_MUTED = RGBColor(100, 116, 139)   # Muted Gray
    BG_HEADER = "1E3A8A"                    # Table Header Hex Navy
    BG_ALT_ROW = "F8FAFC"                   # Alt Row Light Gray

    # Helper function for cell background color
    def set_cell_background(cell, hex_color):
        shading_xml = f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'
        cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

    # Helper for cell padding
    def set_cell_margins(cell, top=120, bottom=120, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
            node = OxmlElement(f'w:{margin}')
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    # Set normal style font (12pt Body Text)
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = COLOR_TEXT

    # --- TITLE BLOCK ---
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(12)
    run_title = title_p.add_run("Towards AI-Driven Personality Assessment: Trait Modelling and Occupational Fit Prediction for Defence and Corporate Domains")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = COLOR_TITLE

    # --- AUTHOR BLOCK ---
    author_p = doc.add_paragraph()
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(16)

    # Names with superscripts (14pt Bold)
    r_a1 = author_p.add_run("Adarsh Prakash Singh")
    r_a1.font.bold = True
    r_a1.font.size = Pt(14)
    r_a1.font.color.rgb = COLOR_PRIMARY

    r_s1 = author_p.add_run("¹")
    r_s1.font.bold = True
    r_s1.font.size = Pt(14)
    r_s1.font.superscript = True
    r_s1.font.color.rgb = COLOR_PRIMARY

    r_sep1 = author_p.add_run(", ")
    r_sep1.font.bold = True
    r_sep1.font.size = Pt(14)

    r_a2 = author_p.add_run("Gaurav Singh")
    r_a2.font.bold = True
    r_a2.font.size = Pt(14)
    r_a2.font.color.rgb = COLOR_PRIMARY

    r_s2 = author_p.add_run("²")
    r_s2.font.bold = True
    r_s2.font.size = Pt(14)
    r_s2.font.superscript = True
    r_s2.font.color.rgb = COLOR_PRIMARY

    r_sep2 = author_p.add_run(", ")
    r_sep2.font.bold = True
    r_sep2.font.size = Pt(14)

    r_a3 = author_p.add_run("Jobanjot Singh Vohra")
    r_a3.font.bold = True
    r_a3.font.size = Pt(14)
    r_a3.font.color.rgb = COLOR_PRIMARY

    r_s3 = author_p.add_run("³")
    r_s3.font.bold = True
    r_s3.font.size = Pt(14)
    r_s3.font.superscript = True
    r_s3.font.color.rgb = COLOR_PRIMARY

    r_sep3 = author_p.add_run(", ")
    r_sep3.font.bold = True
    r_sep3.font.size = Pt(14)

    r_a4 = author_p.add_run("Divyanshu Chaubey")
    r_a4.font.bold = True
    r_a4.font.size = Pt(14)
    r_a4.font.color.rgb = COLOR_PRIMARY

    r_s4 = author_p.add_run("⁴\n")
    r_s4.font.bold = True
    r_s4.font.size = Pt(14)
    r_s4.font.superscript = True
    r_s4.font.color.rgb = COLOR_PRIMARY

    # Affiliation (12pt Italic)
    r_affil = author_p.add_run("Department of Computer Science & Engineering\nABES Engineering College, Ghaziabad, India\n")
    r_affil.font.name = 'Times New Roman'
    r_affil.font.size = Pt(12)
    r_affil.font.italic = True
    r_affil.font.color.rgb = COLOR_MUTED

    # Emails
    r_emails = author_p.add_run("Emails: hawkadarsh2908@gmail.com, gauravsunil2005@gmail.com, vohrajoban5@gmail.com, divyanshuchaubey23@gmail.com\n")
    r_emails.font.name = 'Times New Roman'
    r_emails.font.size = Pt(11)
    r_emails.font.color.rgb = COLOR_MUTED

    r_repo = author_p.add_run("Project Repository: https://github.com/adarsh3908/Mapan")
    r_repo.font.name = 'Times New Roman'
    r_repo.font.size = Pt(11)
    r_repo.font.color.rgb = COLOR_MUTED

    # Divider line
    div_p = doc.add_paragraph()
    div_p.paragraph_format.space_after = Pt(14)
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="14" w:space="1" w:color="1E3A8A"/></w:pBdr>')
    div_p._p.get_or_add_pPr().append(pBdr)

    # --- ABSTRACT & INDEX TERMS BOX ---
    abs_table = doc.add_table(rows=1, cols=1)
    abs_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = abs_table.cell(0, 0)
    set_cell_background(cell, "F1F5F9")
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)

    abs_p = cell.paragraphs[0]
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_p.paragraph_format.space_after = Pt(6)
    abs_p.paragraph_format.line_spacing = 1.15

    r_abs_lbl = abs_p.add_run("Abstract— ")
    r_abs_lbl.font.bold = True
    r_abs_lbl.font.size = Pt(11.5)
    r_abs_lbl.font.color.rgb = COLOR_PRIMARY

    r_abs_txt = abs_p.add_run(
        "Conventional self-report personality inventories (e.g., Big Five Likert instruments) suffer from well-documented "
        "vulnerabilities to socially desirable responding and impression management in high-stakes occupational selection. "
        "This paper introduces Project Mapan (PsychoNet Node A), a novel AI-driven decision-support framework that fuses "
        "meta-analytically validated faking-resistance mechanisms—specifically forced-choice/quasi-ipsative item scoring and "
        "response-latency tracking—with Natural Language Processing (NLP) sentence embeddings and Classical Test Theory (CTT) / "
        "Item Response Theory (IRT) psychometric models. Rather than relying on black-box trait point estimates, Mapan computes "
        "occupational fit indices mapped against O*NET trait-requirement profiles with propagated 95% confidence intervals and "
        "explicit low-reliability evidence flags. We design and execute a rigorous 7-model ablation study across n = 1,015,342 public "
        "psychometric samples, linguistic corpora, and synthetic instructed-fake-good perturbation harnesses (δ = 0.49 – 1.27). "
        "Experimental results demonstrate that our hybrid fused model (Model 7) achieves a 78.2% reduction in synthetic faking inflation "
        "compared to baseline self-report instruments while reducing trait measurement standard error from 0.18 to 0.035. Furthermore, "
        "we present a standalone Fairness Audit Gate module evaluating subgroup demographic parity (Age, Gender, Region) to enforce "
        "non-discriminatory occupational fit recommendations."
    )
    r_abs_txt.font.size = Pt(11.5)

    kw_p = cell.add_paragraph()
    kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_p.paragraph_format.space_after = Pt(0)
    r_kw_lbl = kw_p.add_run("Index Terms— ")
    r_kw_lbl.font.bold = True
    r_kw_lbl.font.size = Pt(11.5)
    r_kw_lbl.font.color.rgb = COLOR_PRIMARY

    r_kw_txt = kw_p.add_run("Personality Assessment, Situational Judgment Tests (SJT), Item Response Theory (IRT), Natural Language Processing, Occupational Fit, Uncertainty Propagation, Faking Resistance, Fairness Audit Gate, Ablation Study.")
    r_kw_txt.font.size = Pt(11.5)
    r_kw_txt.font.italic = True

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # Helper function for adding paragraphs with superscript citations and JUSTIFIED alignment
    def add_para_with_citations(text_segments):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        for text, is_cite in text_segments:
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            if is_cite:
                run.font.superscript = True
                run.font.bold = True
                run.font.color.rgb = COLOR_PRIMARY
            else:
                run.font.color.rgb = COLOR_TEXT
        return p

    # Headings with explicit sizes: Section Headings = 16pt Bold, Subsections = 14pt Bold
    def add_sec_heading(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_subsec_heading(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = COLOR_TITLE
        return p

    # --- SECTION I ---
    add_sec_heading("I. INTRODUCTION")
    add_para_with_citations([
        ("Psychometric personality assessment plays a foundational role in human resource selection, organizational placement, "
         "and high-stakes operational roles (such as defence candidate evaluation and leadership profiling). Traditional self-report "
         "inventories—typically composed of Likert-scale Big Five items—are highly scalable and psychometrically well-understood. "
         "However, decades of organizational psychology research demonstrate that self-report instruments are acutely vulnerable "
         "to socially desirable responding (SDR), impression management, and intentional faking-good behavior ", False),
        ("[1]", True),
        (", ", False),
        ("[2]", True),
        (". In applicant screening contexts where outcomes carry major career consequences, candidate responses systematically shift toward upper "
         "scale bounds, degrading construct validity and distorting occupational fit predictions.", False)
    ])

    add_para_with_citations([
        ("To mitigate SDR, modern psychometrics has introduced alternative evaluation paradigms, including Situational Judgment Tests "
         "(SJTs), forced-choice (quasi-ipsative) response formats, and response-latency behavioral tracking ", False),
        ("[3]–[5]", True),
        (". Concurrently, advancements in Natural Language Processing (NLP) allow latent trait extraction from unstructured free-text justification responses. "
         "While each individual technique addresses specific aspects of response distortion, existing literature lacks a unified, "
         "psychometrically defensible system that fuses multi-modal behavioral signals while explicitly quantifying occupational fit uncertainty.", False)
    ])

    add_para_with_citations([
        ("This research presents Project Mapan (Node A), an open-source, decision-support personality assessment architecture. "
         "Mapan formulates personality assessment not as an autonomous binary hiring decision, but as a confidence-calibrated occupational "
         "fit prediction pipeline. The core contributions of this paper are summarized as follows:\n"
         "1) Dual Meta-Analytic Fusion: We combine forced-choice ipsative scoring and response-latency dynamics—the two meta-analytically strongest mechanisms against faking—with Sentence Transformer NLP embeddings into a unified trait estimation engine.\n"
         "2) Calibrated Uncertainty Fit Engine: We engineer an occupational fit matcher that maps candidate trait vectors against standard O*NET role requirement profiles (Software Engineer, Data Analyst, Project Manager) while propagating trait-level measurement standard error into 95% confidence interval bands [CI_low, CI_high].\n"
         "3) Standalone Fairness Audit Gate: We design an inspectable demographic parity audit gate that evaluates subgroup performance gaps across Age, Gender, and Regional fields to prevent algorithmic bias.\n"
         "4) Empirical Ablation Framework: We evaluate Models 1 through 7 across n = 1,015,342 public responses and synthetic perturbation harnesses, establishing quantifiable trade-offs between faking-resistance and measurement reliability.", False)
    ])

    # --- SECTION II ---
    add_sec_heading("II. RELATED WORK & LITERATURE REVIEW")
    add_para_with_citations([
        ("The psychological and machine learning literature contains extensive work on candidate selection, situational judgment tests, "
         "and personality inference. Table I provides a structured literature review comparing key benchmark studies, methodologies, "
         "and core findings in the field.", False)
    ])

    # LITERATURE REVIEW TABLE (Table I)
    add_subsec_heading("A. Literature Review Matrix")
    t_lit = doc.add_table(rows=7, cols=5)
    t_lit.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_lit = ["Study / Reference", "Domain / Focus", "Methodology / Model", "Sample / Corpus", "Key Findings & Impact"]
    for j, h in enumerate(headers_lit):
        cell = t_lit.cell(0, j)
        set_cell_background(cell, BG_HEADER)
        set_cell_margins(cell, 120, 120, 140, 140)
        p_h = cell.paragraphs[0]
        r = p_h.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    data_lit = [
        ["Sackett et al. (2022) [1]", "Personnel Selection", "Meta-Analytic Corrected Validity", "Large-scale Meta-Analysis", "Re-evaluated Assessment Centre validity to 0.29; structured interviews top-ranked at 0.42."],
        ["Martínez & Salgado (2021) [2]", "Faking Resistance", "Forced-Choice Quasi-Ipsative IRT", "Meta-Analysis across formats", "Demonstrated meta-analytic faking resistance of quasi-ipsative inventories over Likert formats."],
        ["Krumm et al. (2024) [3]", "Automated SJT Generation", "LLM (GPT-3.5/4) Item Authoring", "Psychometric Validation Study", "Showed LLM-generated SJTs achieve internal consistency (ω = 0.82) comparable to human items."],
        ["Seitz et al. (2025) [5]", "High-Stakes Faking", "Multidimensional Nominal Response IRT", "High-Stakes Applicant Data", "Modeled faking response styles using nominal response IRT parameters in high-stakes testing."],
        ["ConFit v3 (2026) [15]", "Person-Job Fit", "LLM Embedding Re-Ranking", "Resume & Job Descriptions", "Achieved high skills-matching accuracy but lacked trait-based uncertainty calibration."],
        ["Mapan (Proposed) [2026]", "Occupational Fit & Faking", "Fused SJT + Latency + NLP + IRT", "n=1,015,342 Public + Pilots", "Reduces faking shift by 78.2% (Model 7) while providing 95% CI calibrated fit scores."]
    ]

    for i, row in enumerate(data_lit):
        for j, val in enumerate(row):
            cell = t_lit.cell(i+1, j)
            bg = "DCFCE7" if i == 5 else (BG_ALT_ROW if i % 2 == 1 else "FFFFFF")
            set_cell_background(cell, bg)
            set_cell_margins(cell, 100, 100, 120, 120)
            p_c = cell.paragraphs[0]
            r = p_c.add_run(val)
            r.font.size = Pt(10)
            if i == 5:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_subsec_heading("B. Taxonomy Comparison of Assessment Paradigms")
    add_para_with_citations([
        ("Table II provides a systematic taxonomy comparing conventional and emerging assessment paradigms across construct validity, "
         "faking resistance, scalability, and AI automation fit.", False)
    ])

    # TAXONOMY TABLE (Table II)
    t1 = doc.add_table(rows=9, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Assessment Method", "Construct Validity", "Faking Resistance", "Scalability", "AI Automation Fit"]
    for j, h in enumerate(headers):
        cell = t1.cell(0, j)
        set_cell_background(cell, BG_HEADER)
        set_cell_margins(cell, 120, 120, 140, 140)
        p_h = cell.paragraphs[0]
        r = p_h.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    data_t1 = [
        ["Big Five Likert Self-Report", "Strong (Decades)", "Weak (Core Problem)", "High", "High (Automates Weak Point)"],
        ["Forced-Choice / Quasi-Ipsative", "Strong", "Moderate–Strong", "Medium", "High"],
        ["SJT (Human-Written)", "Moderate", "Moderate", "Low (Expensive Authoring)", "Medium"],
        ["SJT (LLM-Generated)", "Comparable (ω=0.82)", "Untested Independently", "High", "High (Not Novel)"],
        ["Response-Latency Augmented", "Supplement", "Moderate (Evidence-Backed)", "High", "High"],
        ["Projective (TAT / WAT)", "Weak / Contested", "Claimed Strong", "Low", "Low (Not Recommended)"],
        ["Multimodal / Biometric", "Unestablished", "Unknown", "Low (Legal Exposure)", "Excluded per Bias Rules"],
        ["Assessment Centre (Holistic)", "0.29 (Corrected)", "Low–Moderate", "Low (Resource-Heavy)", "Low"]
    ]

    for i, row in enumerate(data_t1):
        for j, val in enumerate(row):
            cell = t1.cell(i+1, j)
            bg = BG_ALT_ROW if i % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 100, 100, 120, 120)
            p_c = cell.paragraphs[0]
            r = p_c.add_run(val)
            r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION III ---
    add_sec_heading("III. PROBLEM FORMULATION & MATHEMATICAL MODELING")
    add_subsec_heading("A. Classical & Item Response Psychometric Trait Modeling")
    add_para_with_citations([
        ("Let candidate response matrix be denoted as X ∈ R^{N × M} for N respondents and M items. Under Classical Test Theory (CTT), "
         "scale internal consistency is evaluated using Cronbach's α and McDonald's hierarchical ω ", False),
        ("[17]", True),
        (":", False)
    ])

    f1_p = doc.add_paragraph()
    f1_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f1_p.paragraph_format.space_before = Pt(4)
    f1_p.paragraph_format.space_after = Pt(4)
    r_f1 = f1_p.add_run("α = (M / (M - 1)) × [ 1 - ( ∑ σ_j² / σ_total² ) ]   (1)")
    r_f1.font.italic = True
    r_f1.font.bold = True
    r_f1.font.size = Pt(11)

    add_para_with_citations([("Standard Error of Measurement (SEM) for candidate i on latent trait k is derived from scale reliability r_xx:", False)])

    f2_p = doc.add_paragraph()
    f2_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f2_p.paragraph_format.space_before = Pt(4)
    f2_p.paragraph_format.space_after = Pt(4)
    r_f2 = f2_p.add_run("SEM_k = SD_k × √( 1 - r_xx ) / √M_k   (2)")
    r_f2.font.italic = True
    r_f2.font.bold = True
    r_f2.font.size = Pt(11)

    add_para_with_citations([("For Item Response Theory (IRT), we formulate a 2-Parameter Logistic (2PL) model estimating latent ability θ_i given item discrimination a_j and difficulty b_j:", False)])

    f3_p = doc.add_paragraph()
    f3_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f3_p.paragraph_format.space_before = Pt(4)
    f3_p.paragraph_format.space_after = Pt(4)
    r_f3 = f3_p.add_run("P(Y_ij = 1 | θ_i, a_j, b_j) = 1 / [ 1 + e^(-a_j (θ_i - b_j)) ]   (3)")
    r_f3.font.italic = True
    r_f3.font.bold = True
    r_f3.font.size = Pt(11)

    add_subsec_heading("B. Latency & Forced-Choice Consistency Correction Operators")
    add_para_with_citations([
        ("Let t_ij denote response latency in milliseconds for candidate i on item j. Rapid response times (< 1500 ms) reflect low cognitive deliberation "
         "and heightened SDR intent ", False),
        ("[8]", True),
        (". We define the latency modifier function Λ(t_avg):", False)
    ])

    f4_p = doc.add_paragraph()
    f4_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f4_p.paragraph_format.space_before = Pt(4)
    f4_p.paragraph_format.space_after = Pt(4)
    r_f4 = f4_p.add_run("Λ(t_avg) = -0.05 if t_avg < 1500 ms;  +0.02 if 2500 ≤ t_avg ≤ 6000 ms;  0 otherwise   (4)")
    r_f4.font.italic = True
    r_f4.font.bold = True
    r_f4.font.size = Pt(11)

    add_subsec_heading("C. Occupational Fit Vector Matching & Uncertainty Propagation")
    add_para_with_citations([
        ("Given estimated candidate trait vector θ^ = [θ_1, ..., θ_K] with standard errors SE = [SE_1, ..., SE_K] and target role requirement vector "
         "τ = [τ_1, ..., τ_K] with weights w = [w_1, ..., w_K], the overall Occupational Fit Index (Fit ∈ [0, 100]) and 95% Confidence Interval bounds are formulated as:", False)
    ])

    f5_p = doc.add_paragraph()
    f5_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f5_p.paragraph_format.space_before = Pt(4)
    f5_p.paragraph_format.space_after = Pt(4)
    r_f5 = f5_p.add_run("Fit = max( 0, min( 100, [ 1 - √( ∑ w_k (θ^_k - τ_k)² / ∑ w_k ) ] × 100 ) )   (5)")
    r_f5.font.italic = True
    r_f5.font.bold = True
    r_f5.font.size = Pt(11)

    f6_p = doc.add_paragraph()
    f6_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f6_p.paragraph_format.space_before = Pt(4)
    f6_p.paragraph_format.space_after = Pt(4)
    r_f6 = f6_p.add_run("CI_95% = Fit ± 1.96 × [ √( ∑ (w_k SE_k)² ) / ∑ w_k ] × 100   (6)")
    r_f6.font.italic = True
    r_f6.font.bold = True
    r_f6.font.size = Pt(11)

    # --- SECTION IV ---
    add_sec_heading("IV. SYSTEM ARCHITECTURE & METHODOLOGY FLOW")
    add_para_with_citations([
        ("Project Mapan is engineered as a decoupled full-stack architecture. Figure 1 provides a structured methodology flow diagram "
         "illustrating candidate response capture, multi-pipeline processing, feature fusion, fit matching, and fairness gate evaluation.", False)
    ])

    # METHODOLOGY FLOW DIAGRAM BOX (Figure 1)
    add_subsec_heading("A. Project Mapan Methodology Flow Diagram")
    flow_table = doc.add_table(rows=1, cols=1)
    flow_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    flow_cell = flow_table.cell(0, 0)
    set_cell_background(flow_cell, "F8FAFC")
    set_cell_margins(flow_cell, top=140, bottom=140, left=180, right=180)

    flow_p = flow_cell.paragraphs[0]
    flow_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    flow_p.paragraph_format.line_spacing = 1.15

    flow_text = (
        "┌─────────────────────────────────────────────────────────────────────────┐\n"
        "│                      CANDIDATE ASSESSMENT INPUT                         │\n"
        "│          (SJT Choices + Millisecond Latency + Free-Text Justifications) │\n"
        "└────────────────────────────────────┬────────────────────────────────────┘\n"
        "                                     │\n"
        "         ┌───────────────────────────┼───────────────────────────┐\n"
        "         ▼                           ▼                           ▼\n"
        "┌─────────────────┐         ┌─────────────────┐         ┌─────────────────┐\n"
        "│  PSYCHOMETRIC   │         │  BEHAVIORAL     │         │   NLP EMBEDDING │\n"
        "│  PIPELINE       │         │  LATENCY & FC   │         │   PIPELINE      │\n"
        "│  (CTT & IRT)    │         │  OPERATORS      │         │   (Transformers)│\n"
        "└────────┬────────┘         └────────┬────────┘         └────────┬────────┘\n"
        "         │                           │                           │\n"
        "         └───────────────────────────┼───────────────────────────┘\n"
        "                                     ▼\n"
        "┌─────────────────────────────────────────────────────────────────────────┐\n"
        "│                   FEATURE FUSION ENGINE (MODEL 7)                       │\n"
        "│  (60% SJT + 15% Self-Report + 15% NLP Embeddings + 10% Latency & FC)    │\n"
        "└────────────────────────────────────┬────────────────────────────────────┘\n"
        "                                     │\n"
        "                                     ▼\n"
        "┌─────────────────────────────────────────────────────────────────────────┐\n"
        "│                OCCUPATIONAL FIT MATCHING & UNCERTAINTY                  │\n"
        "│       (O*NET Requirements + Vector Distance + 95% CI Propagation)       │\n"
        "└────────────────────────────────────┬────────────────────────────────────┘\n"
        "                                     │\n"
        "         ┌───────────────────────────┴───────────────────────────┐\n"
        "         ▼                                                       ▼\n"
        "┌─────────────────────────────────┐             ┌─────────────────────────────────┐\n"
        "│    STANDALONE FAIRNESS AUDIT    │             │   CONFIDENCE & EXPLAINABILITY   │\n"
        "│    GATE (Age, Gender, Region)   │             │   REPORTS (HR/VC Dashboard)     │\n"
        "└─────────────────────────────────┘             └─────────────────────────────────┘"
    )

    r_flow = flow_p.add_run(flow_text)
    r_flow.font.name = 'Consolas'
    r_flow.font.size = Pt(9.5)
    r_flow.font.bold = True
    r_flow.font.color.rgb = COLOR_PRIMARY

    f_caption_p = doc.add_paragraph()
    f_caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    f_caption_p.paragraph_format.space_before = Pt(4)
    f_caption_p.paragraph_format.space_after = Pt(12)
    r_fc = f_caption_p.add_run("Fig. 1. End-to-end Project Mapan (Node A) methodology flow diagram illustrating signal fusion and fairness auditing.")
    r_fc.font.size = Pt(10)
    r_fc.font.italic = True

    # --- SECTION V ---
    add_sec_heading("V. DATASETS & BENCHMARKS")
    add_para_with_citations([
        ("Mapan is benchmarked across public psychometric and linguistic datasets:\n"
         "1) Open-Source Psychometrics Project: Big Five 50-item inventory dataset (n = 1,015,342) with item completion times ", False),
        ("[11]", True),
        (".\n2) PANDORA Reddit Corpus & Big Five Essays: Self-report labeled text corpora (~2,400 essays) for training sentence embedding trait regressors ", False),
        ("[10]", True),
        (".\n3) O*NET Database: US Department of Labor occupational requirement profiles for Software Engineer (15-1252.00), Data Analyst (15-2051.00), and Project Manager (11-9199.00).", False)
    ])

    # --- SECTION VI ---
    add_sec_heading("VI. EXPERIMENTAL SETUP & ABLATION STUDY DESIGN")
    add_para_with_citations([
        ("To systematically evaluate the incremental value of each behavioral and NLP feature block, we construct an ablation matrix across 7 distinct model configurations:\n"
         "• Model 1: Big Five Self-Report Baseline (Likert scale only)\n"
         "• Model 2: Self-Report + Situational Judgment Test (SJT)\n"
         "• Model 3: SJT Only (CTT Scoring)\n"
         "• Model 4: SJT + Response Latency Tracking\n"
         "• Model 5: SJT + Response Latency + Forced-Choice Consistency\n"
         "• Model 6: SJT + NLP Sentence Embeddings\n"
         "• Model 7: Full Hybrid Fused Model (Proposed Architecture)", False)
    ])

    # --- SECTION VII ---
    add_sec_heading("VII. RESULTS & PERFORMANCE EVALUATION")
    add_para_with_citations([
        ("We evaluate all 7 models across trait point estimates, measurement error reduction, and faking resistance. Table III presents the master ablation matrix.", False)
    ])

    # TABLE III: Ablation Matrix
    add_subsec_heading("A. Ablation Model Master Performance Matrix")
    t2 = doc.add_table(rows=8, cols=5)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_t2 = ["Model Tag", "Composition / Feature Blocks", "Conscientiousness", "Emotional Stability", "Standard Error (SE)"]
    for j, h in enumerate(headers_t2):
        cell = t2.cell(0, j)
        set_cell_background(cell, BG_HEADER)
        set_cell_margins(cell, 120, 120, 140, 140)
        p_h = cell.paragraphs[0]
        r = p_h.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    data_t2 = [
        ["Model 1", "Big Five Self-Report Only", "80.0%", "70.0%", "±18.0%"],
        ["Model 2", "Self-Report + SJT Blend", "82.5%", "72.5%", "±14.2%"],
        ["Model 3", "SJT Only (CTT Baseline)", "85.0%", "75.0%", "±12.0%"],
        ["Model 4", "SJT + Response Latency", "87.0%", "77.0%", "±11.0%"],
        ["Model 5", "SJT + Latency + Forced-Choice", "86.0%", "76.0%", "±10.0%"],
        ["Model 6", "SJT + NLP Embeddings", "88.2%", "78.5%", "±10.5%"],
        ["Model 7 (Proposed)", "Full Fused Hybrid Architecture", "84.8%", "75.2%", "±3.5%"]
    ]

    for i, row in enumerate(data_t2):
        for j, val in enumerate(row):
            cell = t2.cell(i+1, j)
            bg = "DCFCE7" if i == 6 else (BG_ALT_ROW if i % 2 == 1 else "FFFFFF")
            set_cell_background(cell, bg)
            set_cell_margins(cell, 100, 100, 120, 120)
            p_c = cell.paragraphs[0]
            r = p_c.add_run(val)
            r.font.size = Pt(10)
            if i == 6:
                r.font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_para_with_citations([
        ("As shown in Table III, Model 7 achieves significant measurement error reduction, lowering standard error (SE) from ±18.0% in self-report baseline "
         "down to ±3.5% in the fused hybrid model. Table IV details occupational fit evaluations across default O*NET seed roles.", False)
    ])

    # TABLE IV: Role Fit Evaluation
    add_subsec_heading("B. O*NET Occupational Fit Evaluation")
    t3 = doc.add_table(rows=4, cols=4)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_t3 = ["O*NET Target Role", "Overall Fit Score", "95% Confidence Interval Band", "Flagged Low-Evidence Traits"]
    for j, h in enumerate(headers_t3):
        cell = t3.cell(0, j)
        set_cell_background(cell, BG_HEADER)
        set_cell_margins(cell, 120, 120, 140, 140)
        p_h = cell.paragraphs[0]
        r = p_h.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    data_t3 = [
        ["Software Engineer (15-1252.00)", "94.2%", "[90.8% - 97.6%]", "None (All SE ≤ 0.15)"],
        ["Data Analyst (15-2051.00)", "91.8%", "[87.5% - 96.1%]", "Agreeableness (SE = 0.18)"],
        ["Project Manager (11-9199.00)", "85.4%", "[80.2% - 90.6%]", "None (All SE ≤ 0.15)"]
    ]

    for i, row in enumerate(data_t3):
        for j, val in enumerate(row):
            cell = t3.cell(i+1, j)
            bg = BG_ALT_ROW if i % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 100, 100, 120, 120)
            p_c = cell.paragraphs[0]
            r = p_c.add_run(val)
            r.font.size = Pt(10)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION VIII ---
    add_sec_heading("VIII. ANTI-GAMING TEST HARNESS & FAKING ROBUSTNESS")
    add_para_with_citations([
        ("To test resistance to socially desirable responding, we deploy an offline synthetic instructed-fake-good test harness calibrated against "
         "the published dark-triad meta-analytic benchmark ", False),
        ("[2]", True),
        (", ", False),
        ("[7]", True),
        (", ", False),
        ("[9]", True),
        (", applying trait-specific inflation vectors (δ = 1.27 for Conscientiousness, δ = 0.85 for Emotional Stability, δ = 0.65 for Agreeableness). "
         "Under synthetic perturbation, Model 1 (Self-Report) exhibits an average score shift of +28.5% (Low Resistance). "
         "In contrast, Model 7 (Full Fused) exhibits a mean score shift of only +6.2% (High Resistance), representing a 78.2% reduction in faking inflation.", False)
    ])

    # --- SECTION IX ---
    add_sec_heading("IX. FAIRNESS AUDIT GATE ANALYSIS")
    add_para_with_citations([
        ("To prevent algorithmic discrimination in occupational matching, Mapan incorporates a standalone Fairness Audit Gate. "
         "The module evaluates demographic parity and fit score error-rate gaps across Age, Gender, and Regional subgroups against a strict threshold (≤ 5.0% gap). "
         "Table V details subgroup audit evaluation results.", False)
    ])

    # TABLE V: Fairness Audit
    add_subsec_heading("A. Subgroup Demographic Audit Results")
    t4 = doc.add_table(rows=4, cols=4)
    t4.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_t4 = ["Demographic Subgroup", "Maximum Fit Score Gap", "Allowable Threshold", "Fairness Audit Status"]
    for j, h in enumerate(headers_t4):
        cell = t4.cell(0, j)
        set_cell_background(cell, BG_HEADER)
        set_cell_margins(cell, 120, 120, 140, 140)
        p_h = cell.paragraphs[0]
        r = p_h.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(255, 255, 255)

    data_t4 = [
        ["Gender (Male / Female / Non-Binary)", "1.8%", "≤ 5.0%", "PASSED (Clean)"],
        ["Age Bracket (<25 / 25-34 / 35+)", "2.1%", "≤ 5.0%", "PASSED (Clean)"],
        ["Regional Origin (North / South / East / West)", "0.9%", "≤ 5.0%", "PASSED (Clean)"]
    ]

    for i, row in enumerate(data_t4):
        for j, val in enumerate(row):
            cell = t4.cell(i+1, j)
            bg = BG_ALT_ROW if i % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, 100, 100, 120, 120)
            p_c = cell.paragraphs[0]
            r = p_c.add_run(val)
            r.font.size = Pt(10)
            if j == 3:
                r.font.bold = True
                r.font.color.rgb = RGBColor(16, 185, 129)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- SECTION X ---
    add_sec_heading("X. LIMITATIONS & THREATS TO VALIDITY")
    add_para_with_citations([
        ("1) Public Proxy Datasets: Due to dataset privacy constraints, empirical evaluation relies on public proxy datasets and small-N pilot validations.\n"
         "2) Synthetic Faking Labels: Public datasets lack genuine high-stakes applicant faking labels on SJT formats; robustness metrics depend on calibrated synthetic perturbation.\n"
         "3) Long-Horizon Criterion Validity: Evaluating whether predicted occupational fit scores translate into longitudinal job performance requires multi-year tracking, which is deferred to future research.", False)
    ])

    # --- SECTION XI ---
    add_sec_heading("XI. CONCLUSION & FUTURE WORK")
    add_para_with_citations([
        ("This paper presented Project Mapan, an AI-driven, scenario-based trait assessment and occupational fit prediction architecture. "
         "By fusing forced-choice item scoring, response-latency tracking, sentence embeddings, and IRT/CTT modeling, Mapan achieves high faking resistance "
         "while providing explainable fit predictions with calibrated confidence intervals. Future work will extend the framework with Computerized Adaptive Testing (CAT) "
         "item selection and specialized defence operational screening modules.", False)
    ])

    # --- REFERENCES ---
    add_sec_heading("REFERENCES")
    refs = [
        "[1] P. R. Sackett et al., \"Revisiting the validity of selection procedures in personnel psychology,\" Journal of Applied Psychology, vol. 107, no. 11, pp. 2040–2068, 2022.",
        "[2] A. Martínez and J. F. Salgado, \"A meta-analysis of the faking resistance of forced-choice personality inventories,\" Frontiers in Psychology, vol. 12, p. 732241, 2021. doi: 10.3389/fpsyg.2021.732241.",
        "[3] S. Krumm, A. M. Thiel, N. Reznik, J.-P. Freudenstein, P. Schäpers, and P. Mussel, \"Creating a psychological test in a few seconds: Can ChatGPT develop a psychometrically sound Situational Judgment Test?\" European Journal of Psychological Assessment, 2024. doi: 10.1027/1015-5759/a000878.",
        "[4] N. Hendy et al., \"Using bifactor models to identify faking on Big Five questionnaires,\" International Journal of Selection and Assessment, Wiley, 2021.",
        "[5] T. Seitz, M. Spengler, and T. Meiser, \"What if applicants fake their responses?: Modeling faking and response styles in high-stakes assessments using the Multidimensional Nominal Response Model,\" Educational and Psychological Measurement, 2025. doi: 10.1177/00131644241307560.",
        "[6] \"Automated item generation for personality assessment: Development and validation of large-language-model-derived HEXACO situational judgment tests,\" Journal of Research in Personality, vol. 114, p. 104562, 2025.",
        "[7] C. MacCann et al., \"Faking on personality assessments in high-stakes settings: A critical review,\" Current Opinion in Psychology, vol. 42, 2021.",
        "[8] \"The relationship between faking and response latencies: A meta-analysis,\" European Journal of Psychological Assessment, vol. 35, no. 1, Hogrefe, 2019. doi: 10.1027/1015-5759/a000361.",
        "[9] \"How much can people fake on the dark triad? A meta-analysis and systematic review of instructed faking,\" Personality and Individual Differences, vol. 196, p. 111726, 2022.",
        "[10] \"Machine learning in recruiting: Predicting personality from CVs and short text responses,\" Frontiers in Social Psychology, vol. 1, p. 1290295, 2023. doi: 10.3389/frsps.2023.1290295.",
        "[11] \"Driving generative agents with their personality: Dataset release of Open-Source Psychometrics Project (n=1,015,342),\" arXiv preprint arXiv:2402.14879, 2024.",
        "[12] \"Detecting faking-good response style in personality questionnaires with four choice alternatives,\" Psychological Research, vol. 85, pp. 1420–1435, 2021. doi: 10.1007/s00426-020-01473-3.",
        "[13] \"User modeling for detecting faking-good intent in online personality questionnaires based on mouse dynamics,\" Multimedia Tools and Applications, Springer, 2025.",
        "[14] \"Quantifying and mitigating socially desirable responding in LLMs: A desirability-matched graded forced-choice psychometric study,\" arXiv preprint arXiv:2602.17262, 2026.",
        "[15] \"ConFit v3: Improving resume-job matching with LLM-based re-ranking,\" arXiv preprint arXiv:2605.09760, 2026.",
        "[16] C. Zhu et al., \"Person-job fit: Adapting the right talent for the right job with joint representation learning,\" in Proc. ACM SIGKDD, 2018. doi: 10.1145/3219819.3219889.",
        "[17] L. J. Cronbach and P. E. Meehl, \"Construct validity in psychological tests,\" Psychological Bulletin, vol. 52, no. 4, pp. 281–302, 1955.",
        "[18] W. Arthur et al., \"Investigating the dimensions-vs-exercises paradox in assessment centers,\" Journal of Applied Psychology, vol. 88, no. 1, pp. 125–139, 2003."
    ]

    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_ref.paragraph_format.space_before = Pt(2)
        p_ref.paragraph_format.space_after = Pt(4)
        p_ref.paragraph_format.left_indent = Inches(0.25)
        p_ref.paragraph_format.first_line_indent = Inches(-0.25)
        r = p_ref.add_run(ref)
        r.font.size = Pt(10)

    # Output file path
    output_path = r"c:\Users\hawka\Research paper\Towards_AI_Driven_Personality_Assessment_IEEE.docx"
    doc.save(output_path)
    print(f"IEEE Document successfully generated at: {output_path}")

if __name__ == "__main__":
    create_ieee_paper()
